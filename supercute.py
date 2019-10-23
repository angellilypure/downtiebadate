# -*- coding:UTF-8 -*-
#设置间隔的
import time
#抓数据的
import requests
from bs4 import BeautifulSoup
#把数据弄能看的
import numpy as np
import pandas as pd
import docx
import xlrd

#显示程序开始  没什么太大意义  就是装逼的
print("程序开始运行，请稍等")

# 设置回复量阈值  判断输入的回复量是否为int类型，不是的话重新要求输入
while True:
    try:
        clicknum = int(input("请输入要爬取的回复数（超过该回复数的才会被爬取）:"))
        break
    except ValueError:
        print("请输入数字！")
M = clicknum
# get请求模版,末尾{}为占位符，方便发送不同页面的请求
tiebaname = str(input("请输入贴吧名："))
template_url = "https://tieba.baidu.com/f?kw=" + tiebaname + "&ie=utf-8&pn={}"

# 从一页中提取 帖子
def extra_from_one_page(page_lst):
    '''从一页中提取 帖子'''
    # 临时列表保存字典数据，每一个帖子都是一个字典数据
    tmp = []
    for i in page_lst:
        #判断是否超过阈值(这边的“threadlist_rep_num”实际上是页面上显示的回复量)
        if int(i.find(class_='threadlist_rep_num center_text').text) > M:
            dic = {}
            #点击量
            dic['num'] = int(i.find(class_='col2_left j_threadlist_li_left').text)
            #帖子名称
            dic['name'] = i.find(class_='j_th_tit').text
            #帖子地址
            dic['address'] = 'https://tieba.baidu.com' + i.find(class_='j_th_tit').a['href']
            tmp.append(dic)
    return tmp


# 爬取n页的数据
def search_n_pages(n):
    #爬取n页的数据，确认爬取页数
    target = []

    #发起n次的get请求
    for i in range(n):
        #跟踪进度
        print('页数:',i+1)

        #按照浏览贴吧的自然行为，每一页50条
        target_url = template_url.format(50*(i))
        res = requests.get(target_url)

        #转为bs对象
        soup = BeautifulSoup(res.text,'html.parser')

        #获取该页帖子列表
        page_lst = soup.find_all(class_='j_thread_list clearfix')

        #该页信息保存到target里
        target.extend(extra_from_one_page(page_lst))

        #休息2秒再访问
        time.sleep(0.5)
    return target

#遍历url列表，访问并读取数据保存为word
def url_write_word(url_list):
    tmp = '文件生成完成'
    s = 1
    #遍历URL列表
    for i in url_list:
        tmp1 = []
        print('生成文件',s)
        s += 1
        i = str(i)
        #读取网页
        res = requests.get(i)
        soup = BeautifulSoup(res.text,'html.parser')
        url_title = soup.find(class_='core_title_txt pull-left text-overflow').text

        t1 = url_title
        t1 = t1[0:7] + '.docx'
        #创建一个word文档
        doc = docx.Document()
        url_word = soup.find(class_='d_post_content j_d_post_content').text
        t2 = url_word
        #将数据写入word
        doc.add_paragraph(i)
        doc.add_paragraph(t2)
        #保存文档
        doc.save(t1)
    return tmp


#从生成的表格中读取url列表，使url_write_word函数可以有读取用的url_list
def execl_read_url(execl_list):
    #打开指定表格
    rbook = xlrd.open_workbook(execl_list)
    rbook.sheets()
    rsheet = rbook.sheet_by_index(0)
    tmp = []
    #遍历表格每一行
    for row in rsheet.get_rows():
        #保留指定的列数数据
        product_column = row[3]
        product_value = product_column.value
        #去掉标题栏，将剩下的url保存为一个列表
        if product_value != 'address':
            r = str(row[3])
            r1 = r[6:-1]
            tmp.append(r1)

    return tmp

#爬取贴吧前n页数据 如果输入的不是int类型 则要求重新输入
while True:
    try:
        num = int(input("请输入要爬取的页数："))
        break
    except ValueError:
        print("请输入数字！")

d = search_n_pages(num)

# 转化为pandas.DataFrame对象
data = pd.DataFrame(d)

# 导出到excel表格
xlsxname = tiebaname + '.xlsx'
data.to_excel(xlsxname)

#读取表格，生成url列表
execl_list = xlsxname
#将url列表保存到url_list
url_list = execl_read_url(execl_list)
#生成word文档
url_write_word(url_list)


